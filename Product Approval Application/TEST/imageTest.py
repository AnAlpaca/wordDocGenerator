from docx import Document
from docx.shared import Inches


doc = docx.Document('myWordDoc.docx')

tables = doc.tables

tables[0].rows[0].cells[0]._element.clear_content()

#Then when the image is inserted to the cell it is not placed one linefeed down.
img = tables[0].rows[0].cells[0].add_paragraph().add_run().add_picture('Image.png', width=Inches(0.4))
