from __future__ import print_function
import xlrd
import xlwt
import time
import os
import os.path
import re
from mailmerge import MailMerge
from datetime import date
from docx import Document
class DocGen:
    
    def __init__(self, partType):
        cwd = os.getcwd() 
        self.dir_path = cwd
        self.interface_path = self.dir_path + "\\templates\\interface.xlsm"
        self.type = partType
    def show_path(self):
        print(self.interface_path)

    def open_excel(self):
        os.system('start EXCEL.EXE %s' %(self.interface_path))
        input("Please input all your required information into the fields in the Excel Spreasheet.\nOnce completed please save and close EXCEL.\nPress Enter to generate Production Approval from EXCEL information.")
    

    def read_excel_main(self):
        self.partDesc = self.worksheet.cell(6, 1).value
        self.partNum = self.worksheet.cell(7, 1).value
        self.author = self.worksheet.cell(8, 1).value
        self.supplier = self.worksheet.cell(9, 1).value
        self.partName = self.worksheet.cell(10, 1).value
        self.material = self.worksheet.cell(12, 1).value
        self.numColours = int(self.worksheet.cell(13, 1).value)   
        self.machForce = self.worksheet.cell(16, 1).value
        self.barrelCap = self.worksheet.cell(17, 1).value
        self.numTools = int(self.worksheet.cell(18, 1).value)
        self.tools = []
        for i in range(19, 19 + self.numTools):
            self.tools.append(self.worksheet.cell(i, 1).value)
    
    def mailMerge_Main(self):
        self.template_main = "hpc_main.docx"
        #self.dir_path + "\\templates\\hpc_main.docx"
        self.main_doc_name = 'PA - ' + self.partDesc +'.docx'
        mergeMain = MailMerge(self.template_main)
        
        documentVariablesSet = mergeMain.get_merge_fields()
        documentVariablesList = list(documentVariablesSet)
        lengthlist = len(documentVariablesList)
        print('The following variables are used in the Word Templates', documentVariablesList)

        mergeMain.merge(
            PartFullName = self.partDesc,
            PartNumber= self.partNum,
            SupplierName= self.supplier,
            AuthorName = self.author,
            PartShortName = self.partName,
            MachineClampForce = str(int(self.machForce)),
            BarrelCapacity = str(int(self.barrelCap)),
            RevNo = "1",
            IntroductionProduct = "Endless Summer is launching a new PPE half face mask known as the NAUTILUS PPE. Following successful moulding trials, the below component is approved for production subject to the below parameters, procedures and processes."
            )
        print("we got here")
        # define the name of the directory to be created
        path = self.dir_path + "\\output\hpc\\" + self.partDesc + "\\" + self.partDesc

        try:
         os.makedirs(path)
        except OSError:
            print ("Creation of the directory %s already exists or has failed" % path)
        else:
            print ("Successfully created the new directorys %s" % path)
        mergeMain.write(path + "\\" + self.main_doc_name)

    def docx_replace_regex(self, regex , replace):
        regex1 = re.compile(r"your regex")
        replace1 = r"your replace string"
        filename = self.dir_path + "\\hpc_main.docx"
        self.doc = Document(filename)
        

        for p in self.doc.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                            for p in self.doc.paragraphs:
                                if regex.search(p.text):
                                    inline = p.runs
                                    # Loop added to work with runs (strings with same style)
                                    for i in range(len(inline)):
                                        if regex.search(inline[i].text):
                                            text = regex.sub(replace, inline[i].text)
                                            inline[i].text = text
        
        self.doc.save('result1.docx')

    def docx_replace_regex2(self, regex , replace):
        filename = self.dir_path + "\\hpc_main.docx"
        doc = Document(filename)

        Dictionary = {"production": "funny", "find_this_text":"new_text"}
        for i in Dictionary:
            for p in doc.paragraphs:
                if p.text.find(i)>=0:
                    p.text=p.text.replace(i,Dictionary[i])

        for i in Dictionary:
            for p in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if regex.search(p.text):
                                inline = p.runs
                                # Loop added to work with runs (strings with same style)
                                for i in range(len(inline)):
                                    if regex.search(inline[i].text):
                                        text = regex.sub(replace, inline[i].text)
                                        inline[i].text = text
        
        #save changed document
        self.doc.save('test.docx')

    def read_excel_material(self):
        self.colours = []
        self.partNumList = []
        self.masterbatch = []
        self.dosage= []
        if self.worksheet.cell(7, 3).value == 1:
            for i in range(self.numColours):
                self.colours.append(self.worksheet.cell((8 + (5*i)), 3).value)
                self.partNumList.append(self.worksheet.cell((9 + (5*i)), 3).value)
                self.masterbatch.append(self.worksheet.cell((10 + (5*i)), 3).value)
                self.dosage.append(self.worksheet.cell((11 + (5*i)), 3).value)

        elif self.worksheet.cell(7, 3).value == 0:
            for i in range(self.numColours):
                self.colours.append(self.worksheet.cell((8 + (5*i)), 3).value)
                self.masterbatch.append(self.worksheet.cell((10 + (5*i)), 3).value)
                self.dosage.append(int(self.worksheet.cell((11 + (5*i)), 3).value)) 


    def read_excel_quality(self):
        if self.worksheet.cell(9, 5).value == "":
            self.image_splitline = "Image not chosen."
        else:
            self.image_splitline = self.worksheet.cell(9, 5).value

        if self.worksheet.cell(14, 5).value == "":
            self.image_sink = "Image not chosen."
        else:
            self.image_sink = self.worksheet.cell(14, 5).value

        if self.worksheet.cell(19, 5).value == "":
            self.image_gate = "Image not chosen."
        else:
            self.image_gate = self.worksheet.cell(19, 5).value

        if self.worksheet.cell(24, 5).value == "":
            self.image_contamination = "Image not chosen."
        else:
            self.image_contamination = self.worksheet.cell(24, 5).value

        if self.worksheet.cell(29, 5).value == "":
            self.image_flow = "Image not chosen."
        else:
            self.image_flow = self.worksheet.cell(29, 5).value    

    def logic(self):
        n.read_excel_main()
        n.mailMerge_Main()
        if self.worksheet.cell(5, 3).value == 1:
            n.read_excel_material()
        else:
            pass
        if self.worksheet.cell(5, 5).value == 1:
            n.read_excel_quality()
        else:
            pass
        

    def worksheet_excel(self):
        workbook = xlrd.open_workbook(self.interface_path)
        try:
            for i in range(7):
                worksheet = workbook.sheet_by_index(i)
          
                if worksheet.cell(6, 1).value == None:
                    pass 
                else:
                    self.worksheet = worksheet
                    n.logic()
        except Exception:
            pass    
    
   
n = DocGen("HPC")

n.docx_replace_regex2("production", "funny")

